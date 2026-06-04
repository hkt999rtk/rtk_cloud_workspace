from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from report_model import (
    CORE_MESSAGE,
    CURRENT_STATUS_SUMMARY,
    DESIGN_MATERIALS,
    DOCX_PATH,
    FIG_DIR,
    REPORT_DATE,
    REPORT_LANGUAGE,
    SNAPSHOT_TIME_UTC,
    collect_linode_health,
    make_figures,
    rgb,
)

def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: list[float]) -> None:
    total_dxa = sum(int(width * 567) for width in widths_cm)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 567)))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Arial Unicode MS"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti TC")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*rgb(color))


def add_paragraph(doc: Document, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_run_font(r, bold=True, color="17324D" if level <= 2 else "0F766E")


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    r = p.add_run(text)
    set_run_font(r)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)


def add_callout(doc: Document, title: str, body: str, fill: str = "DDF7F3") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [16.2])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, 11.5, True, "17324D")
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_run_font(r2, 10.5, None, "233241")
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.75))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, 9.5, None, "5B6773")


def add_existing_figure(doc: Document, material: dict[str, object], index: int) -> None:
    source_path = Path(material["path"])
    if not source_path.exists():
        add_callout(
            doc,
            f"素材缺失：{material['title']}",
            f"找不到來源檔案：{material['source']}。此項目保留在 appendix，待素材補齊後重新產生報告。",
            "FFF2CC",
        )
        return
    copied_path = FIG_DIR / f"{index:02d}_{source_path.stem}{source_path.suffix.lower()}"
    copied_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source_path, copied_path)
    add_figure(doc, copied_path, str(material["caption"]))


def add_key_value_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "17324D")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, 9.5, True, "FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if i == 0:
                set_cell_shading(cells[i], "F3F6F8")
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i != 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_font(r, 9.2, i == 0, "233241")
    set_table_width(table, widths)
    doc.add_paragraph()


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Realtek Video / IoT Cloud 狀態報告 | 內部草稿")
    set_run_font(r, 8.5, None, "5B6773")


def configure_doc(doc: Document) -> None:
    configure_section(doc.sections[0])
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti TC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in [
        ("Title", 24, "17324D"),
        ("Subtitle", 12, "5B6773"),
        ("Heading 1", 17, "17324D"),
        ("Heading 2", 13, "0F766E"),
        ("Heading 3", 11.5, "233241"),
    ]:
        st = styles[style_name]
        st.font.name = "Arial Unicode MS"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti TC")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor(*rgb(color))
        st.font.bold = style_name.startswith("Heading") or style_name == "Title"
        st.paragraph_format.space_before = Pt(10 if style_name != "Title" else 0)
        st.paragraph_format.space_after = Pt(6)

    for style_name in ["List Bullet", "List Bullet 2", "List Number"]:
        st = styles[style_name]
        st.font.name = "Arial Unicode MS"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti TC")
        st.font.size = Pt(10.2)
        st.paragraph_format.space_after = Pt(3)


def new_page(doc: Document) -> None:
    doc.add_page_break()


def build_doc() -> None:
    figures = make_figures()
    linode_health = collect_linode_health()
    doc = Document()
    configure_doc(doc)

    # Cover
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Realtek Video / IoT Cloud\n狀態報告")
    set_run_font(r, 25, True, "17324D")
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("AmebaPRO 推廣、Cloud 建置、8 月 Loading Test 與商業 KPI 對齊")
    set_run_font(r, 12.5, None, "5B6773")
    doc.add_paragraph()
    add_callout(
        doc,
        "核心管理訊息",
        CORE_MESSAGE,
        "FFF2CC",
    )
    add_key_value_table(
        doc,
        ["面向", "目前狀態", "下一步或風險"],
        CURRENT_STATUS_SUMMARY,
        [3.0, 6.2, 6.4],
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"版本：主管簡報式 Word 初稿｜日期：{REPORT_DATE}｜語言：{REPORT_LANGUAGE}｜用途：內部狀態審閱")
    set_run_font(r, 9.5, None, "5B6773")
    add_figure(doc, figures["product_to_kpi"], "圖 1：從 AmebaPRO module 到商業 KPI 的路徑")
    new_page(doc)

    add_heading(doc, "第一部分：摘要", 1)
    add_callout(
        doc,
        "一頁結論",
        "這個 Cloud 的目的不是單純展示技術，而是補齊 AmebaPRO 與 IoT module 推廣所需的完整解決方案：module、SDK、App、Cloud onboarding、Video、OTA、Admin 管理與可量測的 Loading Test。8 月的 50,000 IoT / 5,000 IoT Video loading test，應被視為技術驗證與商業信心指標。",
        "DDF7F3",
    )

    add_heading(doc, "1. 為什麼要做這個 Cloud", 2)
    for item in [
        "AmebaPRO 推廣缺少一個可展示、可驗證、可讓客戶快速導入的 Video Cloud。",
        "這不是單純做 server，而是補齊「module + SDK + app + cloud + OTA + video + management」的完整產品路徑。",
        "商業意義：Cloud 讓模組銷售從硬體規格競爭，延伸到解決方案競爭。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. 技術與商業 KPI 的邏輯鏈", 2)
    add_figure(doc, figures["kpi_map"], "圖 2：技術成果如何轉化為商業 KPI")
    add_key_value_table(
        doc,
        ["技術成果", "商業效果", "可觀察 KPI"],
        [
            ["Cloud 可部署", "客戶 PoC 可以快速開始", "PoC 啟動時間、部署成功率"],
            ["Device 可上線", "證明 module 可被 cloud 管理", "active devices、online success rate"],
            ["SDK 可整合", "降低 app / firmware 開發門檻", "SDK adoption、integration time"],
            ["OTA 可管理", "支援售後更新與產品生命週期", "OTA success rate、rollback cases"],
            ["Video 可連線", "支撐 AmebaPRO video value proposition", "video setup success、viewer latency"],
            ["Admin 可觀測", "降低支援與維運盲點", "support effort、incident response time"],
        ],
        [3.2, 6.0, 6.4],
    )

    add_heading(doc, "3. 8 月 Loading Test 目標", 2)
    add_figure(doc, figures["load_targets"], "圖 3：8 月 Loading Test 目標")
    add_paragraph(
        doc,
        "IoT 50,000 與 IoT Video 5,000 不只是技術壓測數字，也是在對內確認資源投入、對外建立客戶信心。測試結果應轉化為可溝通的成功率、延遲、錯誤分類與容量邊界。",
    )

    add_heading(doc, "4. 三層次架構", 2)
    add_figure(doc, figures["three_layer"], "圖 4：Root / Brand Cloud / End User 三層架構")
    add_paragraph(
        doc,
        "Root 是 Realtek Platform 的治理層；Brand Cloud 是客戶或品牌租戶；End User 則在租戶下管理多台 devices。客戶可以選擇自建 private cloud，也可以依附 Realtek-operated cloud；business model 與收費方式仍待定案。",
    )

    add_heading(doc, "5. 目前 Linode Fullset Deployment 現況", 2)
    add_figure(doc, figures["linode"], "圖 5：Linode fullset deployment 的主要服務")
    add_callout(
        doc,
        "Linode 的定位與可搬移性",
        "Linode 在本報告中代表較基礎的 VM / infrastructure 服務，不是 AWS-style managed-service stack。差別在於 PostgreSQL、MQ / message queue、broker、reverse proxy、runtime 等服務需要由我們在 VM / service layer 自行架設與管理，而不是直接依賴 AWS-native managed architecture。這個做法增加維運責任，但也降低 vendor lock-in，讓同一套 cloud foundation 未來較容易移動到 AWS、GCP、Azure、阿里雲或其他平台雲。",
        "DCEBFF",
    )
    add_key_value_table(
        doc,
        ["Server / Component", "主要責任"],
        [
            ["Video Cloud", "device runtime、video、WebRTC/TURN、MQTT/EMQX、OTA、telemetry"],
            ["Account Manager", "users、organizations、brand cloud、device registry、provisioning"],
            ["Admin Dashboard", "fleet、provisioning、service health、audit、customer/operator view"],
            ["Frontend / Website", "產品介紹、文件入口、private cloud / commercial messaging"],
            ["PostgreSQL / Object Storage / TLS", "資料持久化、media/firmware artifacts、HTTPS 與 routing"],
        ],
        [4.5, 11.1],
    )

    add_heading(doc, "6. 目前成果與下一步", 2)
    add_key_value_table(
        doc,
        ["已完成 foundation", "下一步"],
        [
            ["Linode fullset deployment", "補齊正式維運、監控、備份還原與 resource owner"],
            ["OTA lifecycle foundation", "用 loading test 與 sample app 驗證實際操作流程"],
            ["Android / iOS SDK 與 reference app 素材", "整理 developer onboarding 與 demo path"],
            ["Pro2 SDK / ASDK artifact delivery", "支援 AmebaPRO firmware / app development"],
            ["Admin / Account / Video Cloud 分工", "將 service health 與 KPI 報表化"],
        ],
        [7.6, 8.0],
    )

    new_page(doc)
    add_heading(doc, "第二部分：Cloud / Product / KPI 細節", 1)

    add_heading(doc, "1. 架構細節", 2)
    add_paragraph(
        doc,
        "系統採三層式租戶架構，避免平台治理、客戶品牌雲與終端使用者裝置管理混在一起。這個分層是未來 private cloud、brand cloud 租用、Realtek-operated cloud 三種模式能共存的基礎。",
    )
    for item in [
        "Account Manager 是 identity、tenant、user、device registry 的 source of truth。",
        "Video Cloud 是 video runtime、device transport、OTA、telemetry 的 source of truth。",
        "Admin 是 dashboard / BFF，用於觀測、操作與彙整，不是帳號或裝置資料的最終來源。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. 從模組銷售到 Cloud KPI 的端到端鏈條", 2)
    add_key_value_table(
        doc,
        ["階段", "技術內容", "商業連動"],
        [
            ["AmebaPRO module", "硬體能力與 video / IoT foundation", "建立 design-in 的入口"],
            ["SDK / App", "Android、iOS、Native、Go、JS 與 sample flow", "降低客戶開發門檻，縮短 PoC"],
            ["Cloud onboarding", "claim、bind、provisioning、activation", "讓 device 可以被 account / brand cloud 管理"],
            ["Video / OTA / telemetry", "connected product lifecycle", "從硬體銷售延伸為可維運服務"],
            ["Admin / service health", "fleet、health、audit、support view", "降低支援成本，提高客戶信任"],
        ],
        [3.2, 6.2, 6.2],
    )

    add_heading(doc, "3. KPI Framework / 指標框架", 2)
    add_key_value_table(
        doc,
        ["KPI 類型", "建議指標", "用途"],
        [
            ["技術 KPI", "deployment readiness、online success rate、API latency、OTA success、video setup success、load scale", "證明系統能運作、能擴充、能被測量"],
            ["產品 KPI", "SDK integration time、reference app completeness、customer onboarding steps、feature demo coverage", "衡量客戶從拿到 module 到完成 PoC 的難度"],
            ["商業 KPI", "AmebaPRO design-in pipeline、PoC-to-project conversion、brand cloud/customer count、active devices、private deployment opportunities", "連接業務成效與產品投資"],
            ["維運 KPI", "incident response time、backup/restore readiness、monitoring coverage、manual support effort", "判斷正式商用需要多少人力與流程"],
        ],
        [3.0, 8.1, 4.5],
    )

    add_heading(doc, "4. Security / Device Trust / 裝置信任", 2)
    add_paragraph(
        doc,
        "PKI/mTLS 是 security detail，不是本報告的商業主標題；但它仍然是 enterprise customer trust 的基礎。可用一句話理解：合法 device 才能安全上雲，避免未授權連線、錯誤綁定與售後追蹤困難。",
    )
    for item in [
        "Device certificate / mTLS：裝置用憑證證明自己，不靠共用密碼。",
        "Certificate renewal / revocation：支援裝置憑證生命週期與風險處理。",
        "KPI 連動：降低 support risk、提升 enterprise customer trust、支援 private cloud 導入。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. API / Cloud Pattern / 雲端設計模式", 2)
    add_paragraph(
        doc,
        "API 設計採用主流 IoT cloud 的共同設計原則，參考 AWS IoT、Azure IoT、GCP IoT 的常見 pattern，但不宣稱與任何一家雲端完全相同。",
    )
    add_key_value_table(
        doc,
        ["Pattern", "本平台對應能力"],
        [
            ["Device identity", "Account Manager registry + Video Cloud runtime identity + device trust"],
            ["Provisioning", "Claim resolve、account bind、cross-service lifecycle command、activation"],
            ["Device state / shadow", "Device state、online/offline、runtime readiness projection"],
            ["Telemetry", "Product telemetry ingestion、admin fleet health、metrics export"],
            ["OTA", "Firmware publish、target、rollout、report、cancel、download"],
            ["RBAC / audit", "Organization roles、platform admin、audit events、service health"],
        ],
        [4.2, 11.4],
    )

    add_heading(doc, "6. Product Features / 產品功能", 2)
    for item in [
        "OTA lifecycle：firmware publish、target、rollout、report、cancel、download。",
        "Burst / load 管理：load profile、concurrency、metrics、p95/p99 latency、error classification。",
        "MQTT / EMQX：IoT device transport。",
        "WebRTC / TURN：video device viewer setup。",
        "Account / brand cloud / user / device registry。",
        "Admin Dashboard：fleet health、provisioning、service health、audit。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. SDK 與 Reference App", 2)
    add_paragraph(
        doc,
        "Android SDK、iOS SDK 是客戶與 app team 的主要入口；Native、Go、JavaScript 可作為其他平台與測試整合補充。Reference app / sample flow 用於展示 scan、claim、provision、online、video、OTA 的完整體驗。",
    )
    add_key_value_table(
        doc,
        ["SDK / Sample", "用途"],
        [
            ["Android / iOS SDK", "主要 app integration 入口，支援客戶 mobile app 團隊快速 PoC"],
            ["Native / Go / JavaScript", "設備、工具、server-side validation 與 web sample 補充"],
            ["Reference App / Mockup", "展示 scan / claim / provisioning / status / OTA / video flow"],
            ["Pro2 SDK / ASDK artifacts", "支援 AmebaPRO firmware / app development 與 release handoff"],
        ],
        [5.0, 10.6],
    )

    add_heading(doc, "8. 使用流程圖", 2)
    add_figure(doc, figures["onboarding"], "圖 6：Device onboarding 與 provisioning flow")
    for item in [
        "Brand Cloud 建立，user 被加入 brand cloud 或 customer tenant。",
        "User 使用 app / SDK 掃描 claim token / QR。",
        "Account Manager resolve claim and bind，建立或匹配 registry device。",
        "Account Manager 發出 provisioning intent，cross-service worker 呼叫 Video Cloud。",
        "Video Cloud activation，device 上線並回報 telemetry / video status。",
        "Admin Dashboard 查看狀態、OTA、health、audit。",
    ]:
        add_number(doc, item)

    add_heading(doc, "9. 8 月 Loading Test 計畫", 2)
    add_key_value_table(
        doc,
        ["測試類型", "目標規模", "核心量測"],
        [
            ["IoT Loading Test", "50,000 devices", "大量連線、telemetry、state update、API latency、error rate"],
            ["IoT Video Loading Test", "5,000 video devices", "video control plane、WebRTC setup、TURN readiness、viewer setup latency"],
        ],
        [4.5, 3.5, 7.6],
    )
    add_paragraph(
        doc,
        "測試結果將成為對內資源投入與對外客戶信心的依據。若測試失敗，報告也要能清楚分類瓶頸：API、broker、database、TURN/WebRTC、resource limit、credential/token 或測試資料準備問題。",
    )

    add_heading(doc, "10. 維護與維運現實", 2)
    add_callout(
        doc,
        "維運現實",
        "工具可以快速建立 cloud，但正式商用維運需要人。需要投入 backend/cloud engineer、DevOps/SRE、QA/load test、security/release owner、customer support，並補齊監控告警、備份還原、升級回滾、incident response、SLA 與成本管理。",
        "FFF2CC",
    )

    new_page(doc)
    add_heading(doc, "第三部分：操作畫面與使用流程", 1)
    add_callout(
        doc,
        "操作圖片補充",
        "本節使用 submodule 內既有設計稿，不重新發明 UI。正文只放主管與跨團隊溝通最需要的畫面；完整素材來源列在 appendix，方便後續週報沿用同一套框架。",
        "DCEBFF",
    )

    add_heading(doc, "1. Admin Customer View 操作畫面", 2)
    add_paragraph(
        doc,
        "Admin dashboard 的 Customer View 設計稿已補齊 daily operations 視角：fleet overview、device detail、firmware/OTA、stream health。這些畫面可以直接支撐對外 demo 與內部驗收，而不是只用抽象架構圖說明 cloud capability。",
    )
    for idx, material in enumerate(DESIGN_MATERIALS[:4], start=7):
        add_existing_figure(doc, material, idx)
        add_paragraph(doc, f"操作重點：{material['purpose']}")

    add_heading(doc, "2. SDK / Sample App 操作流程", 2)
    add_existing_figure(doc, DESIGN_MATERIALS[4], 11)
    add_key_value_table(
        doc,
        ["畫面區塊", "操作說明"],
        [
            ["Environment Setup", "設定 base URL、WebSocket URL、Account Manager URL、device id 與 redacted credentials。"],
            ["Add Device / Provision", "登入、選擇 organization、resolve claim token、啟動 provisioning operation，並顯示 readiness evidence。"],
            ["Device Configuration", "讀寫 camera info/config，讓 SDK demo 能展示 device config 操作，而不是只停留在 connection test。"],
            ["Camera Monitor", "保留 16:9 preview frame，顯示 snapshot、WebRTC signaling helper 與 stream request 狀態。"],
            ["Debug Report", "輸出可複製、已 redacted、明確 pass/skip/fail reason 的驗證報告。"],
        ],
        [4.3, 11.3],
    )

    add_heading(doc, "3. Frontend / Product Architecture 產品架構素材", 2)
    add_existing_figure(doc, DESIGN_MATERIALS[5], 12)
    add_paragraph(
        doc,
        "Frontend 素材可用於對外產品敘事；Admin 與 SDK mockup 則用於操作證據。後續週報應分清楚：architecture visual 用來說明產品定位，operation screenshot 用來證明 capability 已有可展示介面。",
    )

    new_page(doc)
    add_heading(doc, "第四部分：Linode Staging 部署與設定", 1)
    add_callout(
        doc,
        "目前部署狀態",
        f"本節使用公開 endpoint 的只讀 health check 更新狀態。檢查時間：{SNAPSHOT_TIME_UTC}。這是 status report evidence；正式 sign-off 仍以 product-level evidence bundle 為準。",
        "DDF7F3",
    )
    add_key_value_table(
        doc,
        ["Component", "Public endpoint", "Current runtime shape"],
        [
            ["Video Cloud runtime", "https://video-cloud-staging.realtekconnect.com", "video-cloud-staging-edge；nginx 1.30.x public TLS gateway；proxy to Video Cloud API。"],
            ["Account Manager API", "https://account-manager.video-cloud-staging.realtekconnect.com", "dedicated public VM；nginx TLS；app on 127.0.0.1:18081；local PostgreSQL on 127.0.0.1:5432。"],
            ["Admin dashboard", "https://admin.video-cloud-staging.realtekconnect.com", "dedicated public VM；nginx TLS；Dockerized admin app；local SQLite cache。"],
        ],
        [4.2, 5.8, 5.6],
    )
    add_key_value_table(
        doc,
        ["Configuration boundary", "Non-secret setting / behavior"],
        [
            ["Admin upstreams", "ACCOUNT_MANAGER_BASE_URL=https://account-manager.video-cloud-staging.realtekconnect.com；VIDEO_CLOUD_BASE_URL=https://video-cloud-staging.realtekconnect.com"],
            ["TLS / edge", "Public services are reached through domain-based HTTPS. Raw VM IPs and private app ports are not status-report evidence targets."],
            ["Persistence", "Account Manager uses local PostgreSQL; Admin uses local SQLite cache; Video Cloud uses its service-owned runtime storage configuration."],
            ["Secrets", "Linode token、DNS credentials、DB DSN、JWT/auth secrets、MQTT credentials、object storage keys 不寫入報告，只列 configuration category。"],
        ],
        [4.4, 11.2],
    )
    add_key_value_table(
        doc,
        ["Component", "Check", "Result", "Observed"],
        [[row["component"], row["check"], row["result"], row["observed"]] for row in linode_health],
        [3.5, 3.0, 2.0, 7.1],
    )
    add_heading(doc, "尚未達正式商用的缺口（Production-ready gaps）", 2)
    for item in [
        "Video Cloud staging `/version` 目前仍回報 AppVersion=debug；若要對外稱 release staging，需改成明確 release version。",
        "Public frontend / promotion site 在本 milestone 應明確標示部署完成或 SKIP，不應模糊帶過。",
        "正式 production-like sign-off 仍需要 backup/restore evidence：Account Manager PostgreSQL、Video Cloud PostgreSQL/object storage、Admin SQLite、EMQX 與任何啟用的 broker state。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "審閱清單", 1)
    for item in [
        "摘要可在 5 分鐘內看懂，且能回答：為什麼做、做到哪裡、對業務 KPI 有什麼幫助。",
        "Detail 與 repo 現況一致：Linode deployment、OTA、SDK、MQTT/EMQX、Account/Admin/Video Cloud 分工。",
        "技術與商業鏈條清楚：Cloud 的價值不是展示技術，而是提高 AmebaPRO / module 推廣成功率。",
        "操作圖片能支撐 demo flow：Admin overview/devices/OTA/stream health、SDK setup/provision/config/debug report。",
        "Linode deployment / configuration 只列非敏感資訊，且 health check 以本次產生報告時的只讀檢查為準。",
        "不過度承諾：business model、pricing、SLA、production maintain 都標示為待定或需人力投入。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Appendix：素材與來源索引", 1)
    add_key_value_table(
        doc,
        ["素材", "來源", "用途"],
        [[str(m["title"]), str(m["source"]), str(m["purpose"])] for m in DESIGN_MATERIALS],
        [3.7, 5.8, 6.1],
    )
    add_heading(doc, "完整素材目錄", 2)
    for item in [
        "repos/rtk_cloud_admin/docs/assets/webui-design/：Customer View dashboard screenshots。",
        "repos/rtk_cloud_client/docs/mockups/：Android/iOS/WebApp sample flow mockups 與 ops lab screenshots。",
        "repos/rtk_cloud_frontend/static/assets/：Connect+ website hero、architecture、feature visuals。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "內部來源（Sources / Internal References）", 2)
    for item in [
        "rtk_cloud_workspace/docs/private-cloud-deployment.md",
        "rtk_cloud_workspace/docs/linode-staging-deployment-snapshot.md",
        "rtk_cloud_workspace/docs/product-level-evidence.md",
        "rtk_cloud_workspace/docs/business-model.md",
        "rtk_cloud_workspace/docs/account-manager-admin-boundary.md",
        "rtk_cloud_workspace/docs/video-cloud-load-test-roadmap.md",
        "rtk_cloud_admin/docs/webui-customer-view-design.md",
        "rtk_cloud_client/docs/SAMPLE_UI_DESIGN_SYSTEM.md",
        "rtk_cloud_contracts_doc/PRODUCT_ONBOARDING.md",
        "rtk_cloud_contracts_doc/FIRMWARE_CAMPAIGN.md",
        "rtk_cloud_client/docs/SDK_USER_DELIVERY.md",
    ]:
        add_bullet(doc, item)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
